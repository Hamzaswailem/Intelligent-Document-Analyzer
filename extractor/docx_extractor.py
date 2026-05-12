from docx import Document

def extract_text_from_docx(file_path):
    # Load the document
    doc = Document(file_path)
    full_text = []

    # Iterate through paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Combine text with newlines
    return '\n'.join(full_text)

docx_extracted_text = extract_text_from_docx('test.docx')


# Test
#print(extract_text_from_docx('test.docx'))






"""
docx_stats
"""
# from docx import Document

# def get_docx_stats(file_path):
#     doc = Document(file_path)
    
#     # Extract all text from paragraphs
#     full_text = "\n".join([para.text for para in doc.paragraphs])
    
#     # Word count: Split by whitespace
#     words = full_text.split()
#     word_count = len(words)
    
#     # Character count (with spaces)
#     char_count_total = len(full_text)
    
#     # Character count (without spaces/newlines)
#     char_count_no_spaces = len(full_text.replace(" ", "").replace("\n", "").replace("\r", ""))
    
#     return {
#         "words": word_count,
#         "chars_with_spaces": char_count_total,
#         "chars_no_spaces": char_count_no_spaces
#     }

# # Example Usage
# stats = get_docx_stats('test.docx')
# print(f"Words: {stats['words']}")
# print(f"Characters: {stats['chars_with_spaces']}")